# -*- coding: utf-8 -*-
# Part of Odoo AI Assistant. See LICENSE file for full copyright and licensing details.

import base64
import logging
from io import BytesIO

from odoo import models, fields, api, _
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)


class AIDocument(models.Model):
    """AI Document Processing"""
    _name = 'ai.assistant.document'
    _description = 'AI Document for Processing'
    _order = 'create_date desc'
    _rec_name = 'filename'

    filename = fields.Char(
        string='Filename',
        required=True
    )
    
    file = fields.Binary(
        string='File',
        required=True
    )
    
    mimetype = fields.Char(
        string='MIME Type'
    )
    
    document_type = fields.Selection([
        ('invoice', 'Invoice'),
        ('business_card', 'Business Card'),
        ('contract', 'Contract'),
        ('receipt', 'Receipt'),
        ('id_document', 'ID Document'),
        ('other', 'Other'),
    ], string='Document Type', default='other')
    
    target_model = fields.Char(
        string='Target Model',
        help='Odoo model to create from extracted data'
    )
    
    extracted_text = fields.Text(
        string='Extracted Text',
        readonly=True
    )
    
    extracted_data = fields.Json(
        string='Extracted Data',
        readonly=True
    )
    
    state = fields.Selection([
        ('draft', 'Draft'),
        ('processing', 'Processing'),
        ('extracted', 'Extracted'),
        ('imported', 'Imported'),
        ('error', 'Error'),
    ], string='Status', default='draft')
    
    error_message = fields.Text(
        string='Error Message'
    )
    
    created_record_id = fields.Integer(
        string='Created Record ID'
    )
    
    created_record_model = fields.Char(
        string='Created Record Model'
    )
    
    user_id = fields.Many2one(
        'res.users',
        string='User',
        default=lambda self: self.env.user
    )
    
    provider_id = fields.Many2one(
        'ai.assistant.provider',
        string='AI Provider Used'
    )

    def action_process(self):
        """Process the document with AI"""
        self.ensure_one()
        
        self.write({'state': 'processing', 'error_message': False})
        
        try:
            # Step 1: Extract text from document
            extracted_text = self._extract_text()
            
            # Step 2: Use AI to parse and structure data
            extracted_data = self._ai_extract_data(extracted_text)
            
            self.write({
                'extracted_text': extracted_text,
                'extracted_data': extracted_data,
                'state': 'extracted',
            })
            
        except Exception as e:
            _logger.error('Document processing failed: %s', str(e))
            self.write({
                'state': 'error',
                'error_message': str(e),
            })
            raise UserError(_('Document processing failed: %s') % str(e))
    
    def _extract_text(self):
        """Extract text from the document"""
        if not self.file:
            raise UserError(_('No file attached'))
        
        file_data = base64.b64decode(self.file)
        filename = self.filename.lower()
        
        if filename.endswith('.pdf'):
            return self._extract_pdf_text(file_data)
        elif filename.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
            return self._extract_image_text(file_data)
        elif filename.endswith('.docx'):
            return self._extract_docx_text(file_data)
        elif filename.endswith('.txt'):
            return file_data.decode('utf-8', errors='ignore')
        else:
            raise UserError(_('Unsupported file format'))
    
    def _extract_pdf_text(self, file_data):
        """Extract text from PDF"""
        try:
            from pdfminer.high_level import extract_text as pdf_extract
            from io import BytesIO
            
            text = pdf_extract(BytesIO(file_data))
            return text.strip()
        except ImportError:
            raise UserError(_('pdfminer.six not installed. Install with: pip install pdfminer.six'))
    
    def _extract_image_text(self, file_data):
        """Extract text from image using OCR"""
        try:
            import pytesseract
            from PIL import Image
            
            image = Image.open(BytesIO(file_data))
            text = pytesseract.image_to_string(image)
            return text.strip()
        except ImportError:
            # Fallback to AI-based extraction
            return self._ai_image_to_text(file_data)
    
    def _extract_docx_text(self, file_data):
        """Extract text from Word document"""
        try:
            from docx import Document
            
            doc = Document(BytesIO(file_data))
            return '\n'.join([para.text for para in doc.paragraphs])
        except ImportError:
            raise UserError(_('python-docx not installed. Install with: pip install python-docx'))
    
    def _ai_image_to_text(self, file_data):
        """Use AI to describe image content"""
        provider = self.provider_id or self.env['ai.assistant.provider'].get_default_provider()
        
        if not provider:
            raise UserError(_('No AI provider configured for image processing'))
        
        # For now, return placeholder - this requires vision-capable models
        return _("[Image content - AI vision extraction requires GPT-4 Vision or similar]")
    
    def _ai_extract_data(self, text):
        """Use AI to extract structured data from text"""
        provider = self.provider_id or self.env['ai.assistant.provider'].get_default_provider()
        
        if not provider:
            raise UserError(_('No AI provider configured'))
        
        # Build extraction prompt based on document type
        prompt = self._build_extraction_prompt(text)
        
        # Get AI response
        response = provider.generate_response(prompt)
        
        # Parse JSON response
        import json
        try:
            # Try to find JSON in the response
            import re
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            return {'raw_response': response}
        except json.JSONDecodeError:
            return {'raw_response': response}
    
    def _build_extraction_prompt(self, text):
        """Build extraction prompt based on document type"""
        base_prompt = """Extract structured data from the following document text.
Return ONLY a valid JSON object with the extracted information.
Do not include any explanatory text, just the JSON.

Document text:
"""
        
        type_prompts = {
            'invoice': """
Extract: customer_name, invoice_number, invoice_date, due_date, total_amount, currency, line_items (array of: description, quantity, unit_price, total), tax_amount, notes.
""",
            'business_card': """
Extract: name, company, title, email, phone, mobile, website, address.
""",
            'contract': """
Extract: parties (array of names), contract_type, start_date, end_date, value_amount, currency, key_terms (array), signatures (array of names).
""",
            'receipt': """
Extract: vendor_name, receipt_number, date, total_amount, currency, payment_method, items (array of: description, quantity, price).
""",
        }
        
        specific_prompt = type_prompts.get(self.document_type, """
Extract all relevant information: dates, amounts, names, addresses, contact details, and any other important data.
""")
        
        return base_prompt + specific_prompt + f"\n\nText:\n{text}"
    
    def action_import(self):
        """Import extracted data as a record"""
        self.ensure_one()
        
        if self.state != 'extracted':
            raise UserError(_('Document must be processed first'))
        
        if not self.target_model:
            raise UserError(_('Target model not specified'))
        
        if not self.extracted_data:
            raise UserError(_('No extracted data available'))
        
        try:
            # Map extracted data to model fields
            values = self._map_data_to_model()
            
            # Create record
            record = self.env[self.target_model].create(values)
            
            self.write({
                'state': 'imported',
                'created_record_id': record.id,
                'created_record_model': self.target_model,
            })
            
            return {
                'type': 'ir.actions.act_window',
                'name': _('Created Record'),
                'res_model': self.target_model,
                'res_id': record.id,
                'view_mode': 'form',
                'target': 'current',
            }
            
        except Exception as e:
            _logger.error('Record creation failed: %s', str(e))
            raise UserError(_('Failed to create record: %s') % str(e))
    
    def _map_data_to_model(self):
        """Map extracted data to model fields"""
        data = self.extracted_data if isinstance(self.extracted_data, dict) else {}
        
        # Common mappings for different models
        if self.target_model == 'res.partner':
            return {
                'name': data.get('name') or data.get('customer_name') or data.get('vendor_name'),
                'email': data.get('email'),
                'phone': data.get('phone'),
                'mobile': data.get('mobile'),
                'website': data.get('website'),
                'street': data.get('address'),
                'company_type': 'person' if data.get('name') else 'company',
            }
        elif self.target_model == 'account.move':
            return {
                'move_type': 'in_invoice',
                'partner_id': self._find_or_create_partner(data),
                'invoice_date': data.get('invoice_date') or data.get('date'),
                # Add more field mappings
            }
        
        return data
